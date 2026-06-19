import streamlit as st
import difflib
import pdfplumber
from docx import Document

# Force wide layout page configuration to maximize screen space
st.set_page_config(page_title="Perfect Document Diff Matrix", layout="wide")

def process_docx_elements(uploaded_file):
    """Extracts both flat text paragraphs and tables chronologically from a Word file."""
    try:
        doc = Document(uploaded_file)
        elements = []
        
        # Word stores components chronologically in the body element tree
        for element in doc.element.body:
            if element.tag.endswith('p'): # Paragraph Element
                p_text = element.text_content().strip() if hasattr(element, 'text_content') else ""
                if not p_text:
                    # Fallback to standard paragraph lookup if element text mapping varies
                    for p in doc.paragraphs:
                        if p._element == element:
                            p_text = p.text.strip()
                if p_text:
                    elements.append(('text', p_text))
                    
            elif element.tag.endswith('tbl'): # Table Element
                for t in doc.tables:
                    if t._element == element:
                        table_matrix = []
                        for row in t.rows:
                            row_cells = [cell.text.strip() for cell in row.cells]
                            table_matrix.append(row_cells)
                        if table_matrix:
                            elements.append(('table', table_matrix))
        return elements
    except Exception as e:
        st.error(f"Error reading Word structural elements: {str(e)}")
        return []

def process_pdf_elements(uploaded_file):
    """Extracts text lines and tables chronologically from a PDF file using pdfplumber."""
    try:
        elements = []
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                # Get chronological layout objects
                tables = page.extract_tables()
                page_text = page.extract_text(layout=True)
                
                # If tables exist, separate them safely from flat text lines
                if tables:
                    for table in tables:
                        elements.append(('table', table))
                elif page_text:
                    for line in page_text.splitlines():
                        if line.strip():
                            elements.append(('text', line.strip()))
        return elements
    except Exception as e:
        st.error(f"Error reading PDF structural elements: {str(e)}")
        return []

def highlight_words_only(text1, text2):
    """Compares strings word-by-word. Applies clean red/green background tints with NO strikethroughs."""
    if not text1 and not text2:
        return "", False
        
    words1 = str(text1).split() if text1 else []
    words2 = str(text2).split() if text2 else []
    
    sm = difflib.SequenceMatcher(None, words1, words2)
    output_html = []
    is_modified = False
    
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            output_html.append(" ".join(words1[i1:i2]))
        elif tag == 'replace':
            is_modified = True
            if words1[i1:i2]:
                output_html.append(f'<span class="diff-del">{" ".join(words1[i1:i2])}</span>')
            if words2[j1:j2]:
                output_html.append(f'<span class="diff-add">{" ".join(words2[j1:j2])}</span>')
        elif tag == 'delete':
            is_modified = True
            output_html.append(f'<span class="diff-del">{" ".join(words1[i1:i2])}</span>')
        elif tag == 'insert':
            is_modified = True
            output_html.append(f'<span class="diff-add">{" ".join(words2[j1:j2])}</span>')
            
    return " ".join(output_html), is_modified

def render_diff_engine(file1, file2):
    """Compares the documents side by side, rendering everything inside the original layout templates."""
    ext1 = file1.name.split(".")[-1].lower()
    ext2 = file2.name.split(".")[-1].lower()
    
    elements1 = process_docx_elements(file1) if ext1 == "docx" else process_pdf_elements(file1)
    elements2 = process_docx_elements(file2) if ext2 == "docx" else process_pdf_elements(file2)
    
    all_blocks_html = []
    added_count, deleted_count, modified_count = 0, 0, 0
    max_elements = max(len(elements1), len(elements2))
    
    for idx in range(max_elements):
        el1 = elements1[idx] if idx < len(elements1) else (None, None)
        el2 = elements2[idx] if idx < len(elements2) else (None, None)
        
        type1, val1 = el1
        type2, val2 = el2
        
        # Scenario A: Element was completely added
        if val1 is None:
            added_count += 1
            if type2 == 'text':
                html = f'<div class="text-block block-added"><span class="diff-add">{val2}</span></div>'
            else:
                html = '<table class="diff-table block-added">'
                for row in val2:
                    html += "<tr>" + "".join([f'<td><span class="diff-add">{c if c else ""}</span></td>' for c in row]) + "</tr>"
                html += "</table>"
            all_blocks_html.append((html, 'added'))
            continue
            
        # Scenario B: Element was completely deleted
        if val2 is None:
            deleted_count += 1
            if type1 == 'text':
                html = f'<div class="text-block block-deleted"><span class="diff-del">{val1}</span></div>'
            else:
                html = '<table class="diff-table block-deleted">'
                for row in val1:
                    html += "<tr>" + "".join([f'<td><span class="diff-del">{c if c else ""}</span></td>' for c in row]) + "</tr>"
                html += "</table>"
            all_blocks_html.append((html, 'deleted'))
            continue
            
        # Scenario C: Compare side-by-side elements
        if type1 == 'text' or type2 == 'text':
            # Force conversion to strings to check paragraph word adjustments
            str1 = " ".join([" ".join(row) for row in val1]) if type1 == 'table' else str(val1)
            str2 = " ".join([" ".join(row) for row in val2]) if type2 == 'table' else str(val2)
            
            highlighted, was_mod = highlight_words_only(str1, str2)
            if was_mod:
                modified_count += 1
                html = f'<div class="text-block">{highlighted}</div>'
                all_blocks_html.append((html, 'modified'))
            else:
                html = f'<div class="text-block">{str1}</div>'
                all_blocks_html.append((html, 'equal'))
                
        elif type1 == 'table' and type2 == 'table':
            table_html = '<table class="diff-table">'
            table_was_mod = False
            max_rows = max(len(val1), len(val2))
            
            for r_idx in range(max_rows):
                row1 = val1[r_idx] if r_idx < len(val1) else []
                row2 = val2[r_idx] if r_idx < len(val2) else []
                table_html += "<tr>"
                
                max_cells = max(len(row1), len(row2))
                for c_idx in range(max_cells):
                    c1 = row1[c_idx] if c_idx < len(row1) else ""
                    c2 = row2[c_idx] if c_idx < len(row2) else ""
                    
                    highlighted_cell, cell_mod = highlight_words_only(c1, c2)
                    if cell_mod:
                        table_was_mod = True
                    table_html += f"<td>{highlighted_cell}</td>"
                table_html += "</tr>"
            table_html += "</table>"
            
            if table_was_mod:
                modified_count += 1
                all_blocks_html.append((table_html, 'modified'))
            else:
                all_blocks_html.append((table_html, 'equal'))
                
    return all_blocks_html, added_count, deleted_count, modified_count

def main():
    # Global Theme Injection Layer
    st.markdown("""
    <style>
        @import url('https://googleapis.com');

        html, body, [data-testid="stAppViewContainer"], .stMarkdown, p, h1, h2, h3, span {
            font-family: 'Inter', sans-serif !important;
        }
        .workspace-view {
            background-color: #0d1117;
            padding: 24px;
            border-radius: 8px;
            border: 1px solid #30363d;
            overflow-x: auto;
        }
        .text-block {
            font-family: 'JetBrains Mono', monospace !important;
            color: #c9d1d9;
            font-size: 14px;
            line-height: 1.6;
            margin-bottom: 12px;
            white-space: pre-wrap;
        }
        .diff-table {
            width: 100%;
            border-collapse: collapse;
            font-family: 'JetBrains Mono', monospace !important;
            font-size: 13px !important;
            line-height: 1.5 !important;
            color: #c9d1d9;
            background-color: #161b22;
            border: 1px solid #30363d;
            margin-bottom: 24px;
        }
        .diff-table td {
            padding: 10px 14px;
            border: 1px solid #30363d;
            vertical-align: top;
        }
        .block-added { background-color: rgba(46, 160, 67, 0.04) !important; }
        .block-deleted { background-color: rgba(248, 51, 60, 0.04) !important; }

        /* Pure word-level highlights. Absolutely NO strikethroughs */
        .diff-add {
            background-color: rgba(46, 160, 67, 0.25) !important;
            color: #3fb950 !important;
            padding: 2px 6px !important;
            border-radius: 4px !important;
            font-weight: 500 !important;
            text-decoration: none !important;
        }
        .diff-del {
